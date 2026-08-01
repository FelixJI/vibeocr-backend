"""导出服务

支持将 OCR 结果导出为 Markdown、HTML、Word、Excel、纯文本格式。
"""

import base64
import io
import logging
import re
from pathlib import Path

from vibeocr.backend.models.ocr_result import DISCARDED_BLOCK_TYPES, OCRResult
from vibeocr.backend.tables.blocks import table_model_from_block
from vibeocr.backend.tables.projections import table_model_to_grid
from vibeocr.backend.tables.reducer import build_result_projections
from vibeocr.backend.utils.html_tables import tables_from_result
from vibeocr.backend.utils.markdown_converter import HTML_STYLE
from vibeocr.runtime_contracts.contracts.tables import TableModelV1

logger = logging.getLogger(__name__)


class ExportService:
    """OCR 结果导出服务"""

    SUPPORTED_FORMATS = ["markdown", "html", "txt", "docx", "xlsx"]

    @staticmethod
    def export(
        result: OCRResult,
        output_path: Path,
        fmt: str,
    ) -> bool:
        """导出单个结果到文件

        Args:
            result: OCR 结果
            output_path: 输出文件路径（含文件名）
            fmt: 导出格式 (markdown, html, txt, docx, xlsx)

        Returns:
            是否成功
        """
        exporters = {
            "markdown": ExportService._export_markdown,
            "html": ExportService._export_html,
            "txt": ExportService._export_txt,
            "docx": ExportService._export_docx,
            "xlsx": ExportService._export_xlsx,
        }

        exporter = exporters.get(fmt)
        if not exporter:
            logger.error("不支持的导出格式: %s", fmt)
            return False

        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            return exporter(result, output_path)
        except Exception as e:
            logger.error("导出失败 [%s -> %s]: %s", fmt, output_path, e)
            return False

    @staticmethod
    def get_output_filename(source_name: str, fmt: str) -> str:
        """根据源文件名和格式生成输出文件名"""
        stem = Path(source_name).stem
        ext_map = {
            "markdown": ".md",
            "html": ".html",
            "txt": ".txt",
            "docx": ".docx",
            "xlsx": ".xlsx",
        }
        return stem + ext_map.get(fmt, ".txt")

    @staticmethod
    def get_unique_output_path(output_path: Path) -> Path:
        """如果文件已存在，自动追加 _1, _2 ... 后缀避免覆盖"""
        if not output_path.exists():
            return output_path
        stem = output_path.stem
        suffix = output_path.suffix
        parent = output_path.parent
        counter = 1
        while True:
            candidate = parent / f"{stem}_{counter}{suffix}"
            if not candidate.exists():
                return candidate
            counter += 1

    @staticmethod
    def _export_markdown(result: OCRResult, output_path: Path) -> bool:
        """导出为 Markdown"""
        content = result.markdown_text or result.raw_text
        output_path.write_text(content, encoding="utf-8")

        # 保存图片到 images 子目录
        if result.images:
            img_dir = output_path.parent / (output_path.stem + "_images")
            img_dir.mkdir(parents=True, exist_ok=True)
            for name, data in result.images.items():
                if isinstance(data, bytes):
                    dest = img_dir / name
                    dest.parent.mkdir(parents=True, exist_ok=True)
                    dest.write_bytes(data)

        logger.debug("导出 Markdown: %s", output_path)
        return True

    @staticmethod
    def _export_html(result: OCRResult, output_path: Path) -> bool:
        """导出为 HTML（内嵌 base64 图片）"""
        html_body = result.html_text or result.raw_text
        content_list = list(getattr(result, "content_list", []) or [])
        has_structured_table = any(
            isinstance(block, dict)
            and block.get("type") == "table"
            and (
                block.get("table")
                or block.get("table_body")
                or block.get("html")
                or (block.get("source") or {}).get("source_html")
            )
            for block in content_list
        )
        if has_structured_table:
            projections = build_result_projections(
                result,
                include_raw=False,
                include_markdown=False,
            )
            if projections is None:
                raise RuntimeError("structured HTML projection was cancelled")
            html_body = projections[2]

        # 将 markdown 中的图片引用替换为 base64 内嵌
        if result.images:
            for img_name, data in result.images.items():
                if isinstance(data, bytes):
                    b64 = base64.b64encode(data).decode("ascii")
                    ext = Path(img_name).suffix.lstrip(".") or "png"
                    mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
                    data_uri = f"data:{mime};base64,{b64}"
                    # 替换 markdown 图片语法
                    html_body = html_body.replace(f"({img_name})", f"({data_uri})")
                    html_body = html_body.replace(
                        f'src="{img_name}"', f'src="{data_uri}"'
                    )

        full_html = (
            "<!DOCTYPE html>\n<html lang='zh-CN'>\n<head>\n"
            "<meta charset='utf-8'>\n"
            "<meta name='viewport' content='width=device-width, initial-scale=1'>\n"
            f"<title>{output_path.stem}</title>\n"
            f"{HTML_STYLE}\n"
            f"</head>\n<body>\n{html_body}\n</body>\n</html>"
        )

        output_path.write_text(full_html, encoding="utf-8")
        logger.debug("导出 HTML: %s", output_path)
        return True

    @staticmethod
    def _export_txt(result: OCRResult, output_path: Path) -> bool:
        """导出为纯文本"""
        content = result.raw_text or result.markdown_text
        output_path.write_text(content, encoding="utf-8")
        logger.debug("导出纯文本: %s", output_path)
        return True

    @staticmethod
    def _export_docx(result: OCRResult, output_path: Path) -> bool:
        """导出为 Word 文档"""
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Inches, Pt  # type: ignore[import-untyped]

        doc = Document()
        content_list = getattr(result, "content_list", [])
        table_written = False
        written_table_htmls: set[str] = set()

        if content_list:
            for block in content_list:
                block_type = block.get("type", "text")
                text = block.get("text", "")

                if block_type in DISCARDED_BLOCK_TYPES:
                    continue
                if block_type == "title":
                    level = min(block.get("level", 1), 6)
                    doc.add_heading(text, level=level)
                elif block_type == "text":
                    text_level = block.get("text_level")
                    if text_level and 1 <= text_level <= 6:
                        doc.add_heading(text, level=text_level)
                    elif text:
                        doc.add_paragraph(text)
                elif block_type == "table":
                    table_captions = block.get("table_caption") or []
                    if table_captions:
                        doc.add_paragraph(" ".join(table_captions), style="Caption")
                    html = block.get("table_body", "") or block.get("html", "")
                    if block.get("table") or html:
                        written_table_htmls.add(html)
                        table_written = (
                            ExportService._add_table_model_to_docx(
                                doc,
                                ExportService._table_model_from_block(
                                    block,
                                    fallback_table_id=f"table-{len(written_table_htmls)}",
                                ),
                            )
                            or table_written
                        )
                    table_footnotes = block.get("table_footnote") or []
                    for fn in table_footnotes:
                        if fn:
                            p = doc.add_paragraph(fn)
                            for run in p.runs:
                                run.font.size = Pt(9)
                elif block_type in ("image", "figure"):
                    img_path = block.get("img_path", "")
                    caption = (
                        block.get("image_caption") or block.get("chart_caption") or []
                    )
                    images = result.images or {}
                    img_added = False
                    if img_path and img_path in images:
                        data = images[img_path]
                        if isinstance(data, bytes):
                            # python-docx 对损坏/不支持的图片抛多种异常
                            # (UnrecognizedImageError/ValueError/KeyError 等)，
                            # 失败时降级为下方占位段落，故静默忽略
                            try:
                                doc.add_picture(io.BytesIO(data), width=Inches(5))
                                img_added = True
                            except Exception:
                                pass
                    if not img_added:
                        label = " ".join(caption) if caption else text
                        if label:
                            doc.add_paragraph(f"[图片: {label}]")
                elif block_type in (
                    "equation",
                    "interline_equation",
                    "inline_equation",
                ):
                    if text:
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.font.name = "Consolas"
                        run.font.size = Pt(11)
                elif block_type == "list":
                    items = block.get("list_items", [])
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")
                elif block_type == "code":
                    body = block.get("code_body", "")
                    if body:
                        doc.add_paragraph(body, style="No Spacing")

        # 兜底：content_list 无 table 块时，从 html_text/markdown_text/text_blocks/
        # raw_text 提取表格 HTML 补写为 Word 表格。
        if not table_written:
            for html in tables_from_result(result):
                if html in written_table_htmls:
                    continue
                written_table_htmls.add(html)
                ExportService._add_table_model_to_docx(
                    doc,
                    ExportService._table_model_from_block(
                        {"type": "table", "table_body": html},
                        fallback_table_id=f"table-{len(written_table_htmls)}",
                    ),
                )

        if not content_list and not written_table_htmls:
            # 既无结构化块也无表格：写纯文本（清洗裸 HTML 标签）。
            text = ExportService._strip_html_tags(
                result.raw_text or result.markdown_text
            )
            for line in text.split("\n"):
                doc.add_paragraph(line)

        doc.save(str(output_path))
        logger.debug("导出 Word: %s", output_path)
        return True

    @staticmethod
    def _add_table_model_to_docx(doc, table_model: TableModelV1) -> bool:
        """把 canonical 表格写成带原生合并单元格的 docx 表格。

        Returns:
            是否成功写入了至少一个表格（无有效行时返回 False）。
        """
        if table_model.row_count == 0 or table_model.column_count == 0:
            return False

        table = doc.add_table(
            rows=table_model.row_count,
            cols=table_model.column_count,
        )
        table.style = "Table Grid"

        for cell in table_model.cells:
            anchor = table.cell(cell.row, cell.column)
            if cell.rowspan > 1 or cell.colspan > 1:
                anchor = anchor.merge(
                    table.cell(
                        cell.row + cell.rowspan - 1,
                        cell.column + cell.colspan - 1,
                    )
                )
            anchor.text = cell.text
        return True

    @staticmethod
    def _export_xlsx(result: OCRResult, output_path: Path) -> bool:
        """导出为 Excel"""
        from openpyxl import Workbook

        wb = Workbook()
        ws_text = wb.active
        if ws_text is None:
            ws_text = wb.create_sheet("Sheet")
        content_list = getattr(result, "content_list", [])

        table_count = 0
        has_text = False
        written_table_htmls: set[str] = set()

        if content_list:
            for block in content_list:
                block_type = block.get("type", "text")
                text = block.get("text", "")

                if block_type in DISCARDED_BLOCK_TYPES:
                    continue
                if block_type == "table":
                    table_captions = block.get("table_caption") or []
                    if table_captions:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        ws_text.append([f"[表格标题] {' '.join(table_captions)}"])
                    html = block.get("table_body", "") or block.get("html", "")
                    if block.get("table") or html:
                        written_table_htmls.add(html)
                        table_count = ExportService._write_xlsx_table_sheet(
                            wb,
                            ExportService._table_model_from_block(
                                block,
                                fallback_table_id=f"table-{table_count + 1}",
                            ),
                            table_count,
                        )
                    table_footnotes = block.get("table_footnote") or []
                    if table_footnotes:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        ws_text.append([f"[表格脚注] {' '.join(table_footnotes)}"])

                elif block_type == "title" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    ws_text.append([f"[标题] {text}"])

                elif block_type == "text" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    text_level = block.get("text_level")
                    if text_level:
                        ws_text.append([f"{'#' * text_level} {text}"])
                    else:
                        ws_text.append([text])

                elif block_type in ("image", "figure"):
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    caption = (
                        block.get("image_caption") or block.get("chart_caption") or []
                    )
                    label = " ".join(caption) if caption else text
                    if label:
                        ws_text.append([f"[图片: {label}"])

                elif block_type == "equation" and text:
                    if not has_text:
                        has_text = True
                        ws_text.title = "文本汇总"
                    ws_text.append([f"[公式] {text}"])

                elif block_type == "list":
                    items = block.get("list_items", [])
                    if items:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        for item in items:
                            ws_text.append([f"• {item}"])

                elif block_type == "code":
                    body = block.get("code_body", "")
                    if body:
                        if not has_text:
                            has_text = True
                            ws_text.title = "文本汇总"
                        ws_text.append([f"[代码] {body}"])

        # 兜底：content_list 无 table 块时，从 html_text/markdown_text/text_blocks
        # 提取表格 HTML，补写为工作表（前后端分离下表格可能只存活在 html_text）。
        if table_count == 0:
            for html in tables_from_result(result):
                if html in written_table_htmls:
                    continue
                written_table_htmls.add(html)
                table_count = ExportService._write_xlsx_table_sheet(
                    wb,
                    ExportService._table_model_from_block(
                        {"type": "table", "table_body": html},
                        fallback_table_id=f"table-{table_count + 1}",
                    ),
                    table_count,
                )

        if not content_list and table_count == 0:
            # 既无结构化块也无表格：写纯文本（清洗裸 HTML 标签，避免把
            # <table><tr> 当文字写进单元格）。
            ws_text.title = "文本"
            text = ExportService._strip_html_tags(
                result.raw_text or result.markdown_text
            )
            for line in text.split("\n"):
                ws_text.append([line])

        if not has_text and table_count > 0:
            if "Sheet" in wb.sheetnames:
                del wb["Sheet"]

        wb.save(str(output_path))
        logger.debug("导出 Excel: %s", output_path)
        return True

    @staticmethod
    def _write_xlsx_table_sheet(wb, table_model: TableModelV1, table_count: int) -> int:
        """把 canonical 表格写成一个「表格 N」工作表，返回新的 table_count。

        无有效行时计数不变（也不创建空工作表）。
        """
        rows_data = table_model_to_grid(table_model)
        if not rows_data:
            return table_count
        table_count += 1
        ws = wb.create_sheet(title=f"表格 {table_count}")
        for row_idx, row in enumerate(rows_data):
            for col_idx, cell_text in enumerate(row):
                ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_text)
        for min_row, min_col, max_row, max_col in table_model.merged_ranges():
            ws.merge_cells(
                start_row=min_row + 1,
                start_column=min_col + 1,
                end_row=max_row + 1,
                end_column=max_col + 1,
            )
        return table_count

    @staticmethod
    def _table_model_from_block(block: dict, *, fallback_table_id: str) -> TableModelV1:
        if isinstance(block.get("table"), dict):
            return table_model_from_block(
                block,
                fallback_table_id=fallback_table_id,
            )

        html = (
            block.get("table_body")
            or block.get("html")
            or (block.get("source") or {}).get("source_html")
        )
        if isinstance(html, str) and "<table" not in html.lower():
            normalized_block = dict(block)
            normalized_block["table_body"] = f"<table>{html}</table>"
            block = normalized_block
        return table_model_from_block(
            block,
            fallback_table_id=fallback_table_id,
        )

    @staticmethod
    def _strip_html_tags(text: str) -> str:
        """剥离 HTML 标签并规整空白（用于纯文本回退，避免裸标签进单元格）。"""
        stripped = re.sub(r"<[^>]+>", " ", text or "")
        return re.sub(r"[ \t]+", " ", stripped).strip()
