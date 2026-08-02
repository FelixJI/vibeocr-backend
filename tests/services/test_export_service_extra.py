"""export_service 补充测试 — 覆盖 txt/html/markdown 导出、不支持的格式、get_output_filename 等"""

import pytest
from vibeocr.backend.models.ocr_result import OCRResult
from vibeocr.backend.services.export_service import ExportService
from vibeocr.runtime_contracts.contracts.tables import TableCellV1, TableModelV1


def _make_result(**kwargs):
    return OCRResult(
        raw_text=kwargs.get("raw_text", ""),
        markdown_text=kwargs.get("markdown_text", ""),
        html_text=kwargs.get("html_text", ""),
        content_list=kwargs.get("content_list", []),
        images=kwargs.get("images", {}),
    )


def _mixed_table_payload() -> dict:
    return TableModelV1(
        table_id="canonical-table",
        row_count=2,
        column_count=3,
        cells=(
            TableCellV1(cell_id="a", row=0, column=0, rowspan=2, text="A"),
            TableCellV1(cell_id="b", row=0, column=1, colspan=2, text="B"),
            TableCellV1(cell_id="c", row=1, column=1, text="C"),
            TableCellV1(cell_id="d", row=1, column=2, text="D"),
        ),
    ).to_payload()


class TestExportUnsupportedFormat:
    def test_unknown_format_returns_false(self, tmp_path):
        result = _make_result(raw_text="hello")
        out = tmp_path / "test.xyz"
        assert not ExportService.export(result, out, "pdf")


class TestGetOutputFilename:
    def test_markdown(self):
        assert (
            ExportService.get_output_filename("report.pdf", "markdown") == "report.md"
        )

    def test_html(self):
        assert ExportService.get_output_filename("doc.txt", "html") == "doc.html"

    def test_txt(self):
        assert ExportService.get_output_filename("file.docx", "txt") == "file.txt"

    def test_docx(self):
        assert ExportService.get_output_filename("scan.png", "docx") == "scan.docx"

    def test_xlsx(self):
        assert ExportService.get_output_filename("data.pdf", "xlsx") == "data.xlsx"

    def test_unknown_format_falls_back_to_txt(self):
        assert ExportService.get_output_filename("file.pdf", "xyz") == "file.txt"


class TestExportTxt:
    def test_exports_raw_text(self, tmp_path):
        result = _make_result(raw_text="hello world")
        out = tmp_path / "test.txt"
        assert ExportService.export(result, out, "txt")
        assert out.read_text(encoding="utf-8") == "hello world"

    def test_exports_markdown_as_fallback(self, tmp_path):
        result = _make_result(markdown_text="# Title")
        out = tmp_path / "test.txt"
        assert ExportService.export(result, out, "txt")
        assert out.read_text(encoding="utf-8") == "# Title"


class TestExportMarkdown:
    def test_exports_markdown_text(self, tmp_path):
        result = _make_result(markdown_text="## Hello")
        out = tmp_path / "test.md"
        assert ExportService.export(result, out, "markdown")
        assert out.read_text(encoding="utf-8") == "## Hello"

    def test_falls_back_to_raw_text(self, tmp_path):
        result = _make_result(raw_text="plain text")
        out = tmp_path / "test.md"
        assert ExportService.export(result, out, "markdown")
        assert out.read_text(encoding="utf-8") == "plain text"

    def test_saves_images_to_subdir(self, tmp_path):
        import io

        from PIL import Image

        img = Image.new("RGB", (10, 10), "blue")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        result = _make_result(
            markdown_text="![img](images/test.png)",
            images={"images/test.png": img_bytes},
        )
        out = tmp_path / "report.md"
        assert ExportService.export(result, out, "markdown")

        img_file = tmp_path / "report_images" / "images" / "test.png"
        assert img_file.exists()
        assert img_file.read_bytes() == img_bytes


class TestExportHtml:
    def test_structured_html_omits_discarded_blocks(self, tmp_path):
        result = _make_result(
            content_list=[
                {"type": "header", "text": "SECRET_HEADER"},
                {"type": "text", "text": "VISIBLE_BODY"},
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
                {"type": "page_number", "text": "SECRET_PAGE_NUMBER"},
                {"type": "footer", "text": "SECRET_FOOTER"},
            ],
            html_text="<p>LOSSY</p>",
        )
        out = tmp_path / "discarded.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert "VISIBLE_BODY" in content
        assert "SECRET_HEADER" not in content
        assert "SECRET_PAGE_NUMBER" not in content
        assert "SECRET_FOOTER" not in content

    def test_content_list_tables_replace_lossy_html_tables_in_stable_order(
        self, tmp_path
    ):
        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="first",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="first-cell",
                                row=0,
                                column=0,
                                text="CANONICAL_FIRST",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "text", "text": "between"},
                {
                    "type": "table",
                    "table_body": ("<table><tr><td>LEGACY_SECOND</td></tr></table>"),
                },
            ],
            html_text=(
                "<main>"
                "<table><tr><td>LOSSY_FIRST</td></tr></table>"
                "<p>between</p>"
                "<table><tr><td>LOSSY_SECOND</td></tr></table>"
                "</main>"
            ),
        )
        out = tmp_path / "test.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert "LOSSY_FIRST" not in content
        assert "LOSSY_SECOND" not in content
        assert content.index("CANONICAL_FIRST") < content.index("LEGACY_SECOND")
        assert "<p>between</p>" in content

    def test_structured_html_parses_each_table_only_once(self, tmp_path, monkeypatch):
        import vibeocr.backend.tables.reducer as reducer

        original = reducer.table_model_from_block
        calls = 0

        def tracked(block, *args, **kwargs):
            nonlocal calls
            calls += 1
            return original(block, *args, **kwargs)

        monkeypatch.setattr(reducer, "table_model_from_block", tracked)
        result = _make_result(
            content_list=[{"type": "table", "table": _mixed_table_payload()}],
            html_text="<p>LOSSY</p>",
        )

        assert ExportService.export(result, tmp_path / "single-pass.html", "html")
        assert calls == 1

    def test_structured_html_does_not_build_unused_markdown(
        self, tmp_path, monkeypatch
    ):
        import vibeocr.backend.tables.reducer as reducer

        def fail_if_called(*args, **kwargs):
            raise AssertionError("HTML-only export must not build Markdown")

        monkeypatch.setattr(reducer, "table_model_to_markdown", fail_if_called)
        result = _make_result(
            content_list=[{"type": "table", "table": _mixed_table_payload()}],
            html_text="<p>LOSSY</p>",
        )
        output = tmp_path / "html-only.html"

        assert ExportService.export(result, output, "html")
        assert "<table" in output.read_text(encoding="utf-8")

    def test_structured_html_matches_shared_projection_for_rich_blocks(self, tmp_path):
        from vibeocr.backend.tables.reducer import build_result_projections

        result = _make_result(
            content_list=[
                {
                    "type": "title",
                    "text": "Report",
                    "text_level": 2,
                },
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                    "table_caption": "Caption",
                    "table_footnote": "Footnote",
                },
                {"type": "list", "list_items": ["One", "Two"]},
                {"type": "code", "code_body": "print(1)"},
            ],
            html_text="<p>LOSSY</p>",
        )
        projections = build_result_projections(result, include_raw=False)
        assert projections is not None
        expected_html = projections[2]
        output = tmp_path / "shared-projection.html"

        assert ExportService.export(result, output, "html")

        assert expected_html in output.read_text(encoding="utf-8")

    def test_exports_html_text(self, tmp_path):
        result = _make_result(html_text="<p>Hello</p>")
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in content
        assert "<p>Hello</p>" in content

    def test_content_blocks_keep_text_and_table_order_without_html_placeholders(
        self, tmp_path
    ):
        result = _make_result(
            content_list=[
                {"type": "text", "text": "BEFORE"},
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="ordered",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="ordered-cell",
                                row=0,
                                column=0,
                                text="TABLE",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "text", "text": "AFTER"},
            ],
            html_text="<p>LOSSY SUMMARY WITHOUT TABLE PLACEHOLDER</p>",
        )
        out = tmp_path / "ordered.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert content.index("BEFORE") < content.index("TABLE") < content.index("AFTER")
        assert "LOSSY SUMMARY" not in content

    def test_structured_image_table_and_list_export_in_order(self, tmp_path):
        result = _make_result(
            content_list=[
                {
                    "type": "image",
                    "img_path": "fig.png",
                    "image_caption": ["FIGURE"],
                },
                {
                    "type": "table",
                    "table": TableModelV1(
                        table_id="image-order",
                        row_count=1,
                        column_count=1,
                        cells=(
                            TableCellV1(
                                cell_id="image-order-cell",
                                row=0,
                                column=0,
                                text="TABLE",
                            ),
                        ),
                    ).to_payload(),
                },
                {"type": "list", "list_items": ["ONE", "TWO"]},
            ],
            html_text="<p>LOSSY</p>",
            images={"fig.png": b"\x89PNG"},
        )
        out = tmp_path / "mixed.html"

        assert ExportService.export(result, out, "html")

        content = out.read_text(encoding="utf-8")
        assert content.index("FIGURE") < content.index("TABLE") < content.index("ONE")
        assert "data:image/png;base64," in content
        assert "<li>ONE</li>" in content

    def test_embeds_base64_images(self, tmp_path):
        import io

        from PIL import Image

        img = Image.new("RGB", (10, 10), "green")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        result = _make_result(
            html_text='<img src="images/photo.png">',
            images={"images/photo.png": img_bytes},
        )
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "data:image/png;base64," in content

    def test_jpg_mime_type(self, tmp_path):
        result = _make_result(
            html_text='<img src="pic.jpg">',
            images={"pic.jpg": b"\xff\xd8\xff\xe0"},
        )
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "data:image/jpeg;base64," in content


class TestExportDocxExtra:
    def test_legacy_table_uses_native_horizontal_and_vertical_merges(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": (
                        "<table>"
                        '<tr><td rowspan="2">A</td><td colspan="2">B</td></tr>'
                        "<tr><td>C</td><td>D</td></tr>"
                        "</table>"
                    ),
                },
            ],
        )
        out = tmp_path / "test.docx"

        assert ExportService.export(result, out, "docx")

        table = Document(out).tables[0]
        assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
        assert table.cell(0, 1)._tc is table.cell(0, 2)._tc
        assert table.cell(1, 1).text == "C"
        assert table.cell(1, 2).text == "D"

    def test_canonical_table_uses_native_horizontal_and_vertical_merges(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
            ],
        )
        out = tmp_path / "test.docx"

        assert ExportService.export(result, out, "docx")

        table = Document(out).tables[0]
        assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
        assert table.cell(0, 1)._tc is table.cell(0, 2)._tc
        assert table.cell(1, 1).text == "C"
        assert table.cell(1, 2).text == "D"

    def test_title_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "title", "text": "My Title", "level": 2}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        headings = [
            p
            for p in doc.paragraphs
            if (s := p.style) is not None and (s.name or "").startswith("Heading")
        ]
        assert any("My Title" in h.text for h in headings)

    def test_equation_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "equation", "text": "a^2+b^2=c^2"}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        assert any("a^2+b^2=c^2" in p.text for p in doc.paragraphs)

    def test_code_block(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[{"type": "code", "code_body": "print('hi')"}],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        assert any("print" in p.text for p in doc.paragraphs)

    def test_fallback_to_raw_text(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(raw_text="line1\nline2")
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "line1" in texts
        assert "line2" in texts

    def test_table_caption_and_footnote(self, tmp_path):
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<tr><td>X</td></tr>",
                    "table_caption": ["Table 1"],
                    "table_footnote": ["note"],
                },
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Table 1" in all_text
        assert "note" in all_text


class TestExportXlsxExtra:
    def test_legacy_table_preserves_logical_coordinates_and_merges(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": (
                        "<table>"
                        '<tr><td rowspan="2">A</td><td colspan="2">B</td></tr>'
                        "<tr><td>C</td><td>D</td></tr>"
                        "</table>"
                    ),
                },
            ],
        )
        out = tmp_path / "test.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        sheet = workbook["表格 1"]
        assert sheet["A1"].value == "A"
        assert sheet["B1"].value == "B"
        assert sheet["B2"].value == "C"
        assert sheet["C2"].value == "D"
        assert {str(cell_range) for cell_range in sheet.merged_cells.ranges} == {
            "A1:A2",
            "B1:C1",
        }

    def test_canonical_table_preserves_logical_coordinates_and_merges(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table": _mixed_table_payload(),
                },
            ],
        )
        out = tmp_path / "test.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        sheet = workbook["表格 1"]
        assert [
            sheet["A1"].value,
            sheet["B1"].value,
            sheet["B2"].value,
            sheet["C2"].value,
        ] == ["A", "B", "C", "D"]
        assert {str(cell_range) for cell_range in sheet.merged_cells.ranges} == {
            "A1:A2",
            "B1:C1",
        }

    def test_title_block(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[{"type": "title", "text": "Report"}],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        assert ws is not None and ws.title == "文本汇总"

    def test_code_block(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[{"type": "code", "code_body": "x=1"}],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws2 = wb.active
        assert ws2 is not None
        values = [c.value for row in ws2.iter_rows() for c in row if c.value]
        assert any("x=1" in str(v) for v in values)

    def test_fallback_to_raw_text(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(raw_text="hello\nworld")
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws3 = wb.active
        assert ws3 is not None
        values = [c.value for row in ws3.iter_rows() for c in row if c.value]
        assert "hello" in values
        assert "world" in values

    def test_table_only_removes_sheet(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "table", "table_body": "<tr><td>V</td></tr>"},
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        assert "Sheet" not in wb.sheetnames

    def test_table_caption_and_footnote_are_kept_in_summary_order(self, tmp_path):
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<table><tr><td>V</td></tr></table>",
                    "table_caption": ["Caption"],
                    "table_footnote": ["Footnote"],
                }
            ]
        )
        out = tmp_path / "metadata.xlsx"

        assert ExportService.export(result, out, "xlsx")

        workbook = load_workbook(out)
        summary = workbook["文本汇总"]
        values = [row[0] for row in summary.iter_rows(values_only=True)]
        assert values == ["[表格标题] Caption", "[表格脚注] Footnote"]


class TestExportCoverageGaps:
    """覆盖 export_service 中剩余未覆盖分支。"""

    def test_export_returns_false_when_exporter_raises(self, tmp_path, monkeypatch):
        """export() 内层 exporter 抛异常时应捕获并返回 False（lines 60-62）。"""
        result = _make_result(raw_text="hello")
        out = tmp_path / "test.txt"

        def boom(result, output_path):
            raise RuntimeError("boom")

        monkeypatch.setattr(ExportService, "_export_txt", staticmethod(boom))
        assert not ExportService.export(result, out, "txt")

    def test_markdown_skips_non_bytes_image(self, tmp_path):
        """markdown 导出遇到非 bytes 图片应跳过写文件（line 103->102）。"""
        result = _make_result(
            markdown_text="![img](images/x.png)",
            images={"images/x.png": "not-bytes"},
        )
        out = tmp_path / "report.md"
        assert ExportService.export(result, out, "markdown")
        assert out.read_text(encoding="utf-8") == "![img](images/x.png)"

    def test_html_projection_cancelled_raises(self, tmp_path, monkeypatch):
        """build_result_projections 返回 None 时 _export_html 抛 RuntimeError（line 134）。

        直接调用 _export_html（不走 export() 外层 try）以观察异常。
        """
        import vibeocr.backend.services.export_service as export_mod

        monkeypatch.setattr(
            export_mod, "build_result_projections", lambda *a, **k: None
        )
        result = _make_result(
            content_list=[
                {"type": "table", "table_body": "<table><tr><td>X</td></tr></table>"},
            ],
            html_text="<p>x</p>",
        )
        out = tmp_path / "test.html"
        with pytest.raises(RuntimeError):
            ExportService._export_html(result, out)

    def test_docx_skips_discarded_block_types(self, tmp_path):
        """docx 导出应跳过 header/footer 等 discarded 块（line 189）。"""
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {"type": "header", "text": "HDR"},
                {"type": "footer", "text": "FTR"},
                {"type": "page_number", "text": "P1"},
                {"type": "text", "text": "BODY"},
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "BODY" in texts
        assert not any("HDR" in t for t in texts)
        assert not any("FTR" in t for t in texts)

    def test_docx_image_corrupt_bytes_falls_back_to_placeholder(self, tmp_path):
        """docx 图片 bytes 损坏应静默降级为占位段落（lines 235-236）。"""
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {
                    "type": "image",
                    "img_path": "broken.png",
                    "text": "BrokenImg",
                    "image_caption": ["Cap"],
                },
            ],
            images={"broken.png": b"not-an-image"},
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        # 占位段落用 caption
        assert any("Cap" in p.text for p in doc.paragraphs)

    def test_docx_add_table_returns_false_for_empty_model(self, tmp_path):
        """_add_table_model_to_docx 在 row/column 为 0 时返回 False（line 295）。"""
        from docx import Document  # type: ignore[import-not-found]

        doc = Document()
        empty = TableModelV1(
            table_id="t",
            row_count=0,
            column_count=0,
            cells=(),
        )
        assert not ExportService._add_table_model_to_docx(doc, empty)

    def test_xlsx_creates_sheet_when_active_is_none(self, tmp_path, monkeypatch):
        """wb.active 为 None 时应创建 Sheet（line 323）。

        通过清空 Workbook._sheets 让原生 active 返回 None。
        """
        import openpyxl
        from openpyxl import Workbook

        wb_instance = Workbook()
        # 清空 sheets 让 active 走 IndexError → None 分支
        wb_instance._sheets = []
        monkeypatch.setattr(openpyxl, "Workbook", lambda: wb_instance)
        result = _make_result(raw_text="hello")
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

    def test_xlsx_table_footnote_emitted(self, tmp_path):
        """xlsx 表格块带 table_footnote 应写入汇总（lines 358-359）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<table><tr><td>X</td></tr></table>",
                    "table_caption": [],
                    "table_footnote": ["fn1"],
                }
            ]
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb["文本汇总"]
        values = [c.value for row in ws.iter_rows() for c in row if c.value]
        assert any("fn1" in str(v) for v in values)

    def test_xlsx_list_items_emitted(self, tmp_path):
        """xlsx list 块带 list_items 应写入汇总（lines 398-404）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "list", "list_items": ["alpha", "beta"]},
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert any("alpha" in v for v in values)
        assert any("beta" in v for v in values)

    def test_xlsx_write_table_sheet_returns_unchanged_for_empty_grid(
        self, tmp_path, monkeypatch
    ):
        """_write_xlsx_table_sheet 在 grid 为空时返回不变计数（line 458）。"""
        import vibeocr.backend.services.export_service as export_mod
        from openpyxl import Workbook

        empty_model = TableModelV1(
            table_id="t",
            row_count=1,
            column_count=1,
            cells=(TableCellV1(cell_id="c", row=0, column=0, text=""),),
        )
        # monkey table_model_to_grid 返回空
        monkeypatch.setattr(export_mod, "table_model_to_grid", lambda model: [])
        wb = Workbook()
        assert ExportService._write_xlsx_table_sheet(wb, empty_model, 0) == 0

    def test_docx_table_block_with_empty_body_skips_add(self, tmp_path):
        """docx 表格块 table_body/html 均空时应跳过 _add_table（branch 204->213）。"""
        from docx import Document  # type: ignore[import-not-found]

        result = _make_result(
            content_list=[
                {"type": "text", "text": "before"},
                {"type": "table"},  # 无 table_body/html/table
                {"type": "text", "text": "after"},
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        doc = Document(str(out))
        assert len(doc.tables) == 0

    def test_docx_table_block_with_empty_footnote_skipped(self, tmp_path):
        """docx 表格块的 table_footnote 含空字符串应被跳过（branch 215->214）。"""
        result = _make_result(
            content_list=[
                {
                    "type": "table",
                    "table_body": "<table><tr><td>X</td></tr></table>",
                    "table_footnote": [""],
                },
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

    def test_docx_fallback_dedup_skips_already_written_table(self, tmp_path):
        """docx 兜底：fallback 表格 html 与 content_list 已写的重复时应 continue（line 265）。"""
        from docx import Document  # type: ignore[import-not-found]

        same_html = "<table><tr><td>DUP</td></tr></table>"
        result = _make_result(
            content_list=[
                {"type": "table", "table_body": same_html},
            ],
            html_text=same_html,
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")
        # content_list 表已写、html_text 同表去重 → 只有 1 张表
        doc = Document(str(out))
        assert len(doc.tables) == 1

    def test_xlsx_skips_discarded_block_types(self, tmp_path):
        """xlsx 导出应跳过 header/footer 等 discarded 块（line 336）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "header", "text": "HDR"},
                {"type": "footer", "text": "FTR"},
                {"type": "text", "text": "BODY"},
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert "BODY" in values
        assert not any("HDR" in v for v in values)

    def test_docx_empty_text_blocks_skipped(self, tmp_path):
        """docx 各块 text/body/items 为空时应静默跳过对应分支。

        覆盖 branches 197->184, 246->184, 255->184, 257->184。
        """
        result = _make_result(
            content_list=[
                {"type": "text", "text_level": 99},  # text_level 越界 & text 空
                {"type": "equation", "text": ""},  # equation 无 text
                {"type": "list", "list_items": []},  # list 空 items
                {"type": "code", "code_body": ""},  # code 无 body
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

    def test_docx_image_non_bytes_data_skipped(self, tmp_path):
        """docx 图片 data 非 bytes 时应跳过嵌入（branch 228->237 + 239->184）。"""
        result = _make_result(
            content_list=[
                {"type": "image", "img_path": "x.png"},  # 无 caption 无 text
            ],
            images={"x.png": "not-bytes"},
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

    def test_docx_fallback_dedup_actually_continues(self, tmp_path):
        """docx 兜底命中 dedup 分支（line 265 continue）。

        content_list 含一个 table_body='<table></table>' 的 table 块 → 解析出空 model
        （row_count=0）故 _add_table_model_to_docx 返回 False、table_written 保持 False，
        但该 html 字符串已加入 written_table_htmls。html_text 含相同空表 → tables_from_result
        返回同串 → 兜底循环检测到 dedup 命中 continue。
        """
        empty_table_html = "<table></table>"
        result = _make_result(
            content_list=[
                {"type": "table", "table_body": empty_table_html},
            ],
            html_text=empty_table_html,
        )
        out = tmp_path / "test.docx"
        # 导出成功，且因 model 为空、dedup 命中，最终无表格
        assert ExportService.export(result, out, "docx")
        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(out))
        assert len(doc.tables) == 0

    def test_xlsx_has_text_true_then_table_caption(self, tmp_path):
        """xlsx 已 has_text=True 时 table_caption 不再重设 title（branch 340->343）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "text", "text": "first"},  # 先让 has_text=True
                {
                    "type": "table",
                    "table_body": "<table><tr><td>X</td></tr></table>",
                    "table_caption": ["cap"],
                },
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert "first" in values
        assert any("cap" in v for v in values)

    def test_xlsx_table_block_no_body_no_caption_footnote(self, tmp_path):
        """xlsx 表格块 html 为空、无 caption 但有 footnote（branches 345->355）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "text", "text": "pre"},  # has_text=True
                {
                    "type": "table",
                    "table_footnote": ["fn"],
                },  # 无 body/html/table/caption
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert any("fn" in v for v in values)

    def test_xlsx_title_and_image_with_has_text_true(self, tmp_path):
        """xlsx 已 has_text=True 时 title/image 不再重设 title（branches 365->368, 381->384）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "text", "text": "pre"},
                {"type": "title", "text": "T"},  # has_text 已 True
                {
                    "type": "image",
                    "img_path": "p",
                    "image_caption": ["IC"],
                },  # has_text 已 True
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert any("T" in v for v in values)

    def test_xlsx_equation_list_code_empty_paths(self, tmp_path):
        """xlsx equation 无 text、list 空 items、code 无 body 的 False 分支（388->331, 399->331, 406->331）。"""
        result = _make_result(
            content_list=[
                {"type": "equation", "text": ""},  # 跳过
                {"type": "list", "list_items": []},  # 跳过
                {"type": "code", "code_body": ""},  # 跳过
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

    def test_xlsx_equation_list_code_with_has_text_true(self, tmp_path):
        """xlsx 已 has_text=True 时 equation/list/code 不重设 title（392->395, 400->403, 408->331, 409->412）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[
                {"type": "text", "text": "pre"},  # has_text=True
                {"type": "equation", "text": "eq"},  # has_text 已 True
                {"type": "list", "list_items": ["it"]},  # has_text 已 True
                {"type": "code", "code_body": "cb"},  # has_text 已 True
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        assert any("eq" in v for v in values)
        assert any("it" in v for v in values)
        assert any("cb" in v for v in values)

    def test_xlsx_fallback_dedup_continues(self, tmp_path):
        """xlsx 兜底 dedup continue（line 419）。

        content_list 含一个 table_body='<table></table>' 的 table 块 → grid 为空、
        _write_xlsx_table_sheet 返回不变计数（table_count=0），但 html 已加入 set。
        html_text 含相同空表 → tables_from_result 返回同串 → 兜底 dedup 命中 continue。
        """
        empty_table_html = "<table></table>"
        result = _make_result(
            content_list=[
                {"type": "table", "table_body": empty_table_html},
            ],
            html_text=empty_table_html,
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        from openpyxl import load_workbook

        wb = load_workbook(str(out))
        # 无有效表格工作表
        assert not any(n.startswith("表格 ") for n in wb.sheetnames)

    def test_xlsx_keeps_default_sheet_when_no_table_and_only_text(self, tmp_path):
        """xlsx 无表、仅文本且 has_text 时，'Sheet' 默认表保留（branch 441->444 False）。"""
        from openpyxl import load_workbook

        result = _make_result(
            content_list=[{"type": "text", "text": "hi"}],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")
        wb = load_workbook(str(out))
        # 默认 Sheet 被重命名为 文本汇总（仍是 active，不在 sheetnames 里叫 "Sheet"）
        assert "文本汇总" in wb.sheetnames or "Sheet" in wb.sheetnames

    def test_html_skips_non_bytes_image(self, tmp_path):
        """html 导出遇到非 bytes 图片应跳过 base64 内嵌（branch 140->139）。"""
        result = _make_result(
            html_text='<img src="x.png">',
            images={"x.png": "not-bytes"},
        )
        out = tmp_path / "test.html"
        assert ExportService.export(result, out, "html")
        content = out.read_text(encoding="utf-8")
        # 非 bytes 不做替换
        assert "data:image" not in content

    def test_docx_unknown_block_type_falls_through(self, tmp_path):
        """docx 未知 block_type 应跳过所有 elif 回到循环（branch 255->184）。"""
        result = _make_result(
            content_list=[
                {"type": "text", "text": "keep"},
                {"type": "unknown_type", "text": "drop"},
            ],
        )
        out = tmp_path / "test.docx"
        assert ExportService.export(result, out, "docx")

    def test_xlsx_unknown_block_type_falls_through(self, tmp_path):
        """xlsx 未知 block_type 应跳过所有 elif 回到循环（branch 388->331）。"""
        result = _make_result(
            content_list=[
                {"type": "text", "text": "keep"},
                {"type": "unknown_type", "text": "drop"},
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

    def test_xlsx_image_no_label_skipped(self, tmp_path):
        """xlsx image 块无 caption 且无 text 时 label 为空应跳过 append（branch 388->331）。"""
        result = _make_result(
            content_list=[
                {"type": "image", "img_path": "x.png"},  # 无 caption 无 text
            ],
        )
        out = tmp_path / "test.xlsx"
        assert ExportService.export(result, out, "xlsx")

    def test_xlsx_no_default_sheet_when_active_renamed(self, tmp_path):
        """xlsx has_text=False、table_count>0 但默认 Sheet 已被重命名时 del 跳过（441->444）。"""
        # content_list 只有 table 块 → has_text=False、table_count=1
        # 默认 active sheet 重命名发生在 has_text=True 路径，此处 has_text=False
        # 故 'Sheet' 仍在 → del 命中（已被 test_table_only_removes_sheet 覆盖 True 分支）。
        # 为覆盖 False 分支：让 wb 一开始就没有 'Sheet'。
        import openpyxl
        from openpyxl import load_workbook
        from vibeocr.backend.services.export_service import ExportService

        result = _make_result(
            content_list=[
                {"type": "table", "table_body": "<table><tr><td>V</td></tr></table>"},
            ],
        )

        from openpyxl import Workbook

        class _NoSheetWB(Workbook):
            def __init__(self):
                super().__init__()
                # 重命名默认 sheet，使 'Sheet' 不在 sheetnames
                if self._sheets:
                    self._sheets[0].title = "Renamed"

        # openpyxl.Workbook 在函数内 import，patch openpyxl.Workbook
        monkeypatch_target = openpyxl.Workbook
        openpyxl.Workbook = _NoSheetWB
        try:
            out = tmp_path / "test.xlsx"
            assert ExportService.export(result, out, "xlsx")
            wb = load_workbook(str(out))
            assert "表格 1" in wb.sheetnames
        finally:
            openpyxl.Workbook = monkeypatch_target
