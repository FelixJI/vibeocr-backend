"""离线表格语义夹具与发布产物门禁。"""

from __future__ import annotations

import hashlib
import json
import os
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from vibeocr.backend.models import ocr_result_from_payload, ocr_result_to_payload
from vibeocr.backend.services.export_service import ExportService
from vibeocr.backend.tables.html_adapter import table_model_from_html
from vibeocr.runtime_contracts.contracts.tables import TableModelV1

REPO_ROOT = Path(__file__).parents[2]
FIXTURE_ROOT = REPO_ROOT / "tests" / "fixtures" / "table_contract" / "v1"


def _fixture_cases() -> list[dict]:
    manifest = json.loads((FIXTURE_ROOT / "manifest.json").read_text(encoding="utf-8"))
    return [
        json.loads((FIXTURE_ROOT / record["file"]).read_text(encoding="utf-8"))
        for record in manifest["fixtures"]
    ]


class _FixturePipeline:
    def __init__(self, responses: list[object]) -> None:
        self._responses = responses

    def predict(self, **_kwargs):
        return iter(self._responses)


class _FixtureService:
    def __init__(self, responses: list[object]) -> None:
        self._pipeline = _FixturePipeline(responses)

    def get_or_create_pipeline(self, _name: str) -> _FixturePipeline:
        return self._pipeline


def _canonical_cells(table: TableModelV1) -> list[tuple[object, ...]]:
    return [
        (
            cell.row,
            cell.column,
            cell.rowspan,
            cell.colspan,
            cell.text,
            cell.is_header,
        )
        for cell in table.cells
    ]


def _first_difference(expected: object, actual: object, path: str = "$") -> str:
    if type(expected) is not type(actual):
        return (
            f"{path}: type mismatch "
            f"expected={type(expected).__name__}, actual={type(actual).__name__}"
        )
    if isinstance(expected, dict):
        expected_keys = set(expected)
        actual_keys = set(actual)
        if expected_keys != actual_keys:
            return (
                f"{path}: keys mismatch "
                f"expected={sorted(expected_keys)}, actual={sorted(actual_keys)}"
            )
        for key in expected:
            difference = _first_difference(expected[key], actual[key], f"{path}.{key}")
            if difference:
                return difference
        return ""
    if isinstance(expected, list):
        if len(expected) != len(actual):
            return (
                f"{path}: length mismatch "
                f"expected={len(expected)}, actual={len(actual)}"
            )
        for index, expected_item in enumerate(expected):
            difference = _first_difference(
                expected_item, actual[index], f"{path}[{index}]"
            )
            if difference:
                return difference
        return ""
    if expected != actual:
        return f"{path}: expected={expected!r}, actual={actual!r}"
    return ""


def _assert_canonical_equal(
    expected: dict, actual: dict, *, fixture_name: str, report_root: Path
) -> None:
    if expected == actual:
        return
    diagnostic_dir = report_root / fixture_name / "canonical-mismatch"
    diagnostic_dir.mkdir(parents=True, exist_ok=True)
    (diagnostic_dir / "expected.json").write_text(
        json.dumps(expected, ensure_ascii=False, indent=2, sort_keys=True),
        encoding="utf-8",
    )
    (diagnostic_dir / "actual.json").write_text(
        json.dumps(actual, ensure_ascii=False, indent=2, sort_keys=True),
        encoding="utf-8",
    )
    difference = _first_difference(expected, actual)
    (diagnostic_dir / "first-difference.txt").write_text(
        difference + "\n", encoding="utf-8"
    )
    pytest.fail(f"{fixture_name}: canonical mismatch: {difference}")


def _result_from_provider_fixture(fixture: dict):
    """Run the actual provider-facing seam for every fixture source."""

    provider = fixture["name"]
    payload = fixture["source"]["provider_payload"]
    if provider.startswith("mineru-"):
        # MinerU response -> OCRResult is the service's established adapter seam.
        from vibeocr.backend.services.mineru_service import MinerUService

        # 与既有 MinerU service contract tests 一致：只实例化转换器，不执行
        # Singleton 的 API 进程启动副作用，保证此门禁完全离线。
        service = MinerUService.__new__(MinerUService)
        return service._build_ocr_result(payload, "fixture.pdf", data=None)
    if provider == "paddleocr-table-recognition":
        from vibeocr.backend.core.pipelines.pipeline_table import (
            TABLE_RECOGNITION_SPEC,
            TableRecognitionOptions,
        )

        return TABLE_RECOGNITION_SPEC.recognize(
            _FixtureService([payload]), None, TableRecognitionOptions()
        )
    if provider == "paddleocr-pp-structure":
        from vibeocr.backend.core.pipelines.pipeline_pp_structure import (
            PP_STRUCTURE_V3_SPEC,
            PPStructureV3Options,
        )

        response = dict(payload)
        response["parsing_res_list"] = [
            SimpleNamespace(**block) for block in payload["parsing_res_list"]
        ]
        return PP_STRUCTURE_V3_SPEC.recognize(
            _FixtureService([response]), None, PPStructureV3Options()
        )
    if provider == "paddleocr-vl":
        from vibeocr.backend.core.pipelines.pipeline_paddlocr_vl import (
            PADDLEOCR_VL_SPEC,
            PaddleOCRVLOptions,
        )

        return PADDLEOCR_VL_SPEC.recognize(
            _FixtureService([payload]), None, PaddleOCRVLOptions()
        )
    raise AssertionError(f"unhandled provider fixture: {provider}")


def test_fixture_manifest_has_stable_hashes_and_only_synthetic_inputs() -> None:
    manifest = json.loads((FIXTURE_ROOT / "manifest.json").read_text(encoding="utf-8"))

    assert manifest["synthetic"] is True
    assert len(manifest["fixtures"]) == 5
    assert {record["provider_version"] for record in manifest["fixtures"]} == {
        "synthetic"
    }
    for record in manifest["fixtures"]:
        fixture = FIXTURE_ROOT / record["file"]
        assert hashlib.sha256(fixture.read_bytes()).hexdigest() == record["sha256"]


def test_canonical_mismatch_writes_expected_actual_and_first_difference(
    tmp_path: Path,
):
    expected = {"schema_version": 1, "cells": [{"text": "expected"}]}
    actual = {"schema_version": 1, "cells": [{"text": "actual"}]}

    with pytest.raises(pytest.fail.Exception, match=r"\$\.cells\[0\]\.text"):
        _assert_canonical_equal(
            expected,
            actual,
            fixture_name="diagnostic-fixture",
            report_root=tmp_path,
        )

    diagnostic_dir = tmp_path / "diagnostic-fixture" / "canonical-mismatch"
    assert (
        json.loads((diagnostic_dir / "expected.json").read_text(encoding="utf-8"))
        == expected
    )
    assert (
        json.loads((diagnostic_dir / "actual.json").read_text(encoding="utf-8"))
        == actual
    )
    assert "$.cells[0].text" in (diagnostic_dir / "first-difference.txt").read_text(
        encoding="utf-8"
    )


@pytest.mark.parametrize("fixture", _fixture_cases(), ids=lambda value: value["name"])
def test_every_provider_fixture_projects_to_the_same_canonical_mixed_span_table(
    fixture,
):
    expected = TableModelV1.from_payload(fixture["expected"]["canonical_table"])
    result = _result_from_provider_fixture(fixture)
    actual_blocks = [
        block
        for block in result.content_list
        if block.get("type") == "table" and isinstance(block.get("table"), dict)
    ]

    assert len(actual_blocks) == 1
    actual = TableModelV1.from_payload(actual_blocks[0]["table"])
    _assert_canonical_equal(
        expected.to_payload(),
        actual.to_payload(),
        fixture_name=fixture["name"],
        report_root=Path(
            os.environ.get("TABLE_CONTRACT_REPORT_DIR", "reports/table-contract")
        ),
    )
    assert actual_blocks[0]["block_id"] == expected.table_id

    repeated = _result_from_provider_fixture(fixture)
    repeated_block = next(
        block
        for block in repeated.content_list
        if block.get("type") == "table" and isinstance(block.get("table"), dict)
    )
    assert repeated.content_list.index(repeated_block) == 0
    assert repeated_block == actual_blocks[0]


@pytest.mark.parametrize("fixture", _fixture_cases(), ids=lambda value: value["name"])
def test_provider_fixture_survives_wire_block_and_three_export_formats(
    fixture: dict, tmp_path: Path
):
    expected = TableModelV1.from_payload(fixture["expected"]["canonical_table"])
    result = _result_from_provider_fixture(fixture)
    wire_result = ocr_result_from_payload(ocr_result_to_payload(result))
    table_block = next(
        block
        for block in wire_result.content_list
        if block.get("type") == "table" and isinstance(block.get("table"), dict)
    )
    _assert_canonical_equal(
        expected.to_payload(),
        TableModelV1.from_payload(table_block["table"]).to_payload(),
        fixture_name=fixture["name"],
        report_root=Path(
            os.environ.get("TABLE_CONTRACT_REPORT_DIR", "reports/table-contract")
        ),
    )
    assert table_block["block_id"] == expected.table_id

    artifact_dir = (
        Path(os.environ.get("TABLE_CONTRACT_REPORT_DIR", tmp_path)) / fixture["name"]
    )
    artifact_dir.mkdir(parents=True, exist_ok=True)
    html_path = artifact_dir / "table.html"
    xlsx_path = artifact_dir / "table.xlsx"
    docx_path = artifact_dir / "table.docx"
    assert ExportService.export(wire_result, html_path, "html")
    assert ExportService.export(wire_result, xlsx_path, "xlsx")
    assert ExportService.export(wire_result, docx_path, "docx")

    html_actual = table_model_from_html(
        html_path.read_text(encoding="utf-8"), table_id=expected.table_id
    )
    assert _canonical_cells(html_actual) == _canonical_cells(expected)

    from openpyxl import load_workbook

    workbook = load_workbook(xlsx_path)
    sheet = workbook["表格 1"]
    assert {str(cell_range) for cell_range in sheet.merged_cells.ranges} == {
        "A1:A2",
        "B1:C1",
    }
    assert sheet["A1"].value == "纵向"
    assert sheet["B1"].value == "横向"

    from docx import Document

    document = Document(docx_path)
    assert len(document.tables) == 1
    document_table = document.tables[0]
    assert document_table.cell(0, 0).text == "纵向"
    assert document_table.cell(0, 1).text == "横向"
    assert document_table.cell(0, 0)._tc is document_table.cell(1, 0)._tc
    assert document_table.cell(0, 1)._tc is document_table.cell(0, 2)._tc
    with zipfile.ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert '<w:gridSpan w:val="2"' in document_xml
    assert '<w:vMerge w:val="restart"' in document_xml
    assert "<w:vMerge/>" in document_xml
